#!/usr/bin/env python3
"""
ingest.py — Index all your PM documents into the local vector store.

Usage:
    python ingest.py              # Smart incremental sync (new/changed files only)
    python ingest.py --reset      # Wipe DB and re-index everything
    python ingest.py --local-only # Skip Google Drive
    python ingest.py --silent     # No output (designed for cron jobs)
"""

import os
import sys
import json
import hashlib
import argparse
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional, Tuple

import yaml
from dotenv import load_dotenv
from rich.console import Console
from rich.panel import Panel
from tqdm import tqdm

load_dotenv()

console = Console()
SILENT = False  # Set to True via --silent flag

def cprint(*args, **kwargs):
    """Print only when not in silent mode."""
    if not SILENT:
        console.print(*args, **kwargs)

# ─── Sync State ───────────────────────────────────────────────────────────────
# .sync_state.json persists between runs to track what's changed

SYNC_STATE_FILE = ".sync_state.json"

def load_sync_state() -> dict:
    """Load persisted sync state (last run timestamp + per-file hashes)."""
    if Path(SYNC_STATE_FILE).exists():
        try:
            with open(SYNC_STATE_FILE) as f:
                return json.load(f)
        except Exception:
            pass
    return {"last_sync": None, "local_hashes": {}, "drive_modtimes": {}}

def save_sync_state(state: dict):
    """Write updated sync state after a successful run."""
    state["last_sync"] = datetime.now(timezone.utc).isoformat()
    with open(SYNC_STATE_FILE, "w") as f:
        json.dump(state, f, indent=2)

# ─── Config ───────────────────────────────────────────────────────────────────

# def load_config() -> dict:
#     with open("config.yaml") as f:
#         return yaml.safe_load(f)
def load_config() -> dict:
    with open("config.yaml") as f:
        config = yaml.safe_load(f)

    # Override with local private config if it exists
    local_cfg_path = Path("config.local.yaml")
    if local_cfg_path.exists():
        with open(local_cfg_path) as f:
            local_cfg = yaml.safe_load(f)
        config.update(local_cfg)

    return config

# ─── Document Loaders ─────────────────────────────────────────────────────────

def load_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def load_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def load_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(v) for v in row if v is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)

def load_text(path: str) -> str:
    try:
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""

def load_file(path: str) -> Optional[str]:
    """Load any supported file type and return its text content."""
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            return load_pdf(path)
        elif ext in (".docx", ".doc"):
            return load_docx(path)
        elif ext in (".xlsx", ".xls"):
            return load_xlsx(path)
        elif ext in (".md", ".txt"):
            return load_text(path)
    except Exception as e:
        cprint(f"  [yellow]⚠ Skipped {Path(path).name}: {e}[/yellow]")
    return None

# ─── File Discovery ───────────────────────────────────────────────────────────

def file_hash(path: str) -> str:
    """MD5 hash of file contents — detects actual changes, not just timestamps."""
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def discover_local_files(config: dict) -> List[str]:
    import fnmatch
    patterns = config.get("file_types", ["*.pdf", "*.docx", "*.md"])
    excludes = config.get("exclude_patterns", [])
    found = []

    for folder_cfg in config.get("local_folders", []):
        base = Path(folder_cfg["path"]).expanduser()
        recursive = folder_cfg.get("recursive", True)
        if not base.exists():
            cprint(f"  [yellow]Folder not found, skipping: {base}[/yellow]")
            continue
        glob_fn = base.rglob if recursive else base.glob
        for pattern in patterns:
            for file_path in glob_fn(pattern):
                skip = any(fnmatch.fnmatch(file_path.name, ex) for ex in excludes)
                if not skip and file_path.is_file():
                    found.append(str(file_path))

    return list(set(found))

# ─── Google Drive ─────────────────────────────────────────────────────────────

def get_drive_service(drive_cfg: dict):
    """OAuth2 auth — opens browser on first run, uses token.json after that."""
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build

    SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
    creds = None
    token_file = drive_cfg.get("token_file", "token.json")
    creds_file = drive_cfg.get("credentials_file", "credentials.json")

    if Path(token_file).exists():
        creds = Credentials.from_authorized_user_file(token_file, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(creds_file, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_file, "w") as f:
            f.write(creds.to_json())

    return build("drive", "v3", credentials=creds)

def fetch_google_drive_docs(config: dict, since: Optional[str] = None) -> List[dict]:
    """
    Fetch Drive files. If `since` is an ISO timestamp, only files modified
    after that time are fetched — this is the delta sync mechanism.
    """
    drive_cfg = config.get("google_drive", {})
    if not drive_cfg.get("enabled", False):
        return []

    try:
        import io
        from googleapiclient.http import MediaIoBaseDownload

        service = get_drive_service(drive_cfg)

        MIME_EXPORT_MAP = {
            "application/vnd.google-apps.document": "text/plain",
            "application/vnd.google-apps.spreadsheet": "text/csv",
        }
        DIRECT_TYPES = {
            "application/pdf", "text/plain", "text/markdown",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }

        all_mimes = list(MIME_EXPORT_MAP.keys()) + list(DIRECT_TYPES)
        mime_filter = " or ".join(f"mimeType = '{m}'" for m in all_mimes)
        query = f"trashed = false and ({mime_filter})"

        # Filter by folder IDs if specified (and not 'root' which implies all)
        folder_ids = [f.get("id") for f in drive_cfg.get("folders", []) if f.get("id") and f.get("id") != "root"]
        if folder_ids:
            folder_filter = " or ".join(f"'{fid}' in parents" for fid in folder_ids)
            query += f" and ({folder_filter})"

        # ── KEY: Tell Drive API to only return files changed since last sync ──
        if since:
            query += f" and modifiedTime > '{since}'"
            cprint(f"  [dim]Fetching Drive files changed since {since[:19]} UTC[/dim]")
        else:
            cprint("  [dim]First run — fetching all Drive files[/dim]")

        docs = []
        page_token = None

        while True:
            resp = service.files().list(
                q=query,
                fields="nextPageToken, files(id, name, mimeType, modifiedTime)",
                pageSize=200,
                pageToken=page_token,
            ).execute()

            for f in resp.get("files", []):
                try:
                    mime = f["mimeType"]
                    if mime in MIME_EXPORT_MAP:
                        data = service.files().export_media(
                            fileId=f["id"], mimeType=MIME_EXPORT_MAP[mime]
                        ).execute()
                        content = data.decode("utf-8", errors="ignore")
                    else:
                        req = service.files().get_media(fileId=f["id"])
                        buf = io.BytesIO()
                        downloader = MediaIoBaseDownload(buf, req)
                        done = False
                        while not done:
                            _, done = downloader.next_chunk()
                        tmp = Path(f"/tmp/pmrag_{f['name']}")
                        tmp.write_bytes(buf.getvalue())
                        content = load_file(str(tmp)) or ""
                        tmp.unlink(missing_ok=True)

                    if content.strip():
                        docs.append({
                            "source": f"GoogleDrive:{f['name']}",
                            "content": content,
                            "modifiedTime": f.get("modifiedTime", ""),
                            "fileId": f["id"],
                        })
                except Exception as e:
                    cprint(f"  [yellow]⚠ Skipped Drive file {f['name']}: {e}[/yellow]")

            page_token = resp.get("nextPageToken")
            if not page_token:
                break

        return docs

    except Exception as e:
        cprint(f"[red]Google Drive error: {e}[/red]")
        return []

# ─── Chunking ─────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    if not text or not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start:start + chunk_size])
        start += chunk_size - overlap
    return chunks

# ─── Vector Store ─────────────────────────────────────────────────────────────

def get_vector_store(config: dict, reset: bool = False):
    import chromadb
    from chromadb.utils import embedding_functions

    db_path = config["vector_store"]["path"]
    if reset and Path(db_path).exists():
        import shutil
        shutil.rmtree(db_path)
        cprint("[yellow]Vector store wiped for full re-index.[/yellow]")

    ef = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=config["embeddings"]["model"]
    )
    client = chromadb.PersistentClient(path=db_path)
    return client.get_or_create_collection(
        name="pm_docs",
        embedding_function=ef,
        metadata={"hnsw:space": "cosine"},
    )

# ─── Core Ingest (also imported by query.py for --sync) ───────────────────────

def run_ingest(config: dict, reset: bool = False, local_only: bool = False) -> Tuple[int, int, int]:
    """
    Incremental ingest. Only processes files that are new or have changed.
    Returns (new_chunks_added, files_skipped, total_chunks_in_db).
    """
    retrieval = config.get("retrieval", {})
    chunk_size = retrieval.get("chunk_size", 800)
    chunk_overlap = retrieval.get("chunk_overlap", 100)

    state = load_sync_state()
    last_sync: Optional[str] = None if reset else state.get("last_sync")
    local_hashes: dict = {} if reset else state.get("local_hashes", {})
    drive_modtimes: dict = {} if reset else state.get("drive_modtimes", {})

    collection = get_vector_store(config, reset=reset)
    added = 0
    skipped = 0

    # ── Local: compare content hashes to find changed files ──
    cprint("\n[bold]📂 Scanning local folders...[/bold]")
    all_local = discover_local_files(config)

    changed_local = []
    for fp in all_local:
        try:
            current_hash = file_hash(fp)
        except Exception:
            continue
        if local_hashes.get(fp) == current_hash:
            skipped += 1
        else:
            changed_local.append((fp, current_hash))

    cprint(f"  Total found: [green]{len(all_local)}[/green] | "
           f"Changed/new: [cyan]{len(changed_local)}[/cyan] | "
           f"Unchanged: [dim]{skipped}[/dim]")

    if changed_local:
        cprint("\n[bold]📄 Indexing changed local files...[/bold]")
        for file_path, current_hash in tqdm(changed_local, desc="Local files", disable=SILENT):
            content = load_file(file_path)
            if not content or not content.strip():
                local_hashes[file_path] = current_hash
                continue

            chunks = chunk_text(content, chunk_size, chunk_overlap)
            doc_id_prefix = f"local:{current_hash}"
            ids = [f"{doc_id_prefix}:{i}" for i in range(len(chunks))]
            metas = [{"source": file_path, "filename": Path(file_path).name, "chunk": i}
                     for i in range(len(chunks))]

            collection.upsert(ids=ids, documents=chunks, metadatas=metas)
            local_hashes[file_path] = current_hash
            added += len(chunks)

    # ── Google Drive: pass last_sync timestamp so Drive API filters for us ──
    if not local_only:
        cprint("\n[bold]☁️  Checking Google Drive...[/bold]")
        drive_docs = fetch_google_drive_docs(config, since=last_sync)

        if drive_docs:
            cprint(f"  [cyan]{len(drive_docs)}[/cyan] new/updated Drive files to index")
            for doc in tqdm(drive_docs, desc="Drive files", disable=SILENT):
                source = doc["source"]
                mod_time = doc.get("modifiedTime", "")
                file_id = doc.get("fileId", source)

                # Unique ID based on file identity + modification time
                h = hashlib.md5(f"{file_id}:{mod_time}".encode()).hexdigest()
                doc_id_prefix = f"drive:{h}"

                chunks = chunk_text(doc["content"], chunk_size, chunk_overlap)
                ids = [f"{doc_id_prefix}:{i}" for i in range(len(chunks))]
                metas = [{"source": source, "filename": source, "chunk": i, "modifiedTime": mod_time}
                         for i in range(len(chunks))]

                collection.upsert(ids=ids, documents=chunks, metadatas=metas)
                drive_modtimes[source] = mod_time
                added += len(chunks)
        else:
            cprint("  [dim]No Drive changes since last sync ✓[/dim]")

    # ── Persist state so next run knows what changed ──
    state["local_hashes"] = local_hashes
    state["drive_modtimes"] = drive_modtimes
    save_sync_state(state)

    return added, skipped, collection.count()

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Ingest PM documents into RAG")
    parser.add_argument("--reset", action="store_true", help="Wipe DB and re-index everything")
    parser.add_argument("--local-only", action="store_true", help="Skip Google Drive")
    parser.add_argument("--silent", action="store_true", help="No output (for cron jobs)")
    args = parser.parse_args()

    global SILENT
    SILENT = args.silent

    if not SILENT:
        console.print(Panel.fit("[bold cyan]PM RAG — Document Ingestion[/bold cyan]", border_style="cyan"))
        state = load_sync_state()
        if state.get("last_sync") and not args.reset:
            last = state["last_sync"][:19].replace("T", " ")
            console.print(f"[dim]Last sync: {last} UTC — only new/changed files will be processed[/dim]")
        elif args.reset:
            console.print("[yellow]--reset: full re-index from scratch[/yellow]")
        else:
            console.print("[dim]First run — indexing all files[/dim]")

    config = load_config()
    added, skipped, total = run_ingest(config, reset=args.reset, local_only=args.local_only)

    if not SILENT:
        console.print(Panel(
            f"[green]✅ Sync complete![/green]\n"
            f"  New chunks indexed:        [cyan]{added}[/cyan]\n"
            f"  Files skipped (unchanged): [yellow]{skipped}[/yellow]\n"
            f"  Total chunks in DB:        [bold]{total}[/bold]",
            title="Done",
            border_style="green",
        ))
        console.print("\n[bold]Run [cyan]python query.py[/cyan] to start asking questions![/bold]")

if __name__ == "__main__":
    main()
